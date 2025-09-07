# app/tools/exporters.py
"""
Export helpers for the Hiring Assistant.

Functions
---------
- checklist_json_to_docx(plan: dict) -> bytes
    Build a single "Hiring Plan" .docx that contains:
    - Target timeline (and, if present, budget + location policy)
    - Checklist (tasks)
    - Interview loop
    - Roles & JDs (mission, requirements, nice-to-haves, benefits)

- jd_to_docx(jd: dict|pydantic_model) -> bytes
    Build a per-role JD .docx (one role per file).

- jds_to_zip(jds_map: dict[str, dict|pydantic_model]) -> bytes
    Bundle multiple JDs into a single ZIP (convenient one-click download).
"""

from io import BytesIO
from typing import Any, Dict
import re
import zipfile

from docx import Document
from docx.shared import Pt


# ----------------------------------------------------------------------
# Internal helpers
# ----------------------------------------------------------------------

def _norm_list(val) -> list[str]:
    """
    Normalize a list-like field that might be:
      - a proper list of strings, **or**
      - a single string with items separated by newlines / commas / semicolons / bullet chars.

    Returns a clean `list[str]`.
    """
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    if isinstance(val, str):
        parts = re.split(r"[\n,;•·\-\u2022]+", val)
        return [p.strip(" •\t\r-") for p in parts if p and p.strip(" •\t\r-")]
    return []


def _safe_name(s: str) -> str:
    """
    Create a filesystem-safe base name for files (e.g., "Founding Engineer" -> "Founding-Engineer").
    Keeps letters, numbers, dash, underscore, dot, and space; then replaces spaces with hyphens.
    """
    base = re.sub(r"[^A-Za-z0-9\-_. ]+", "", (s or "role")).strip().replace(" ", "-")
    return base or "role"


def _add_bullets(doc: Document, items: list[str], style: str = "List Bullet"):
    """
    Add a list of strings as bulleted paragraphs using the given Word style.
    """
    for item in items:
        doc.add_paragraph(item, style=style)


# ----------------------------------------------------------------------
# Combined plan export
# ----------------------------------------------------------------------

def checklist_json_to_docx(plan: Dict[str, Any]) -> bytes:
    """
    Build a clean .docx from the structured checklist JSON.
    Returns raw bytes so Streamlit can offer a direct download.

    Expected keys in `plan` (best-effort; all optional):
      - "timeline_weeks": int
      - "budget_usd": int
      - "location_policy": str
      - "tasks": [ { "name": str, "owner": str, "due": str }, ... ]
      - "interview_loop": [ { "stage": str, "duration_min": int, "signals": [str,...] }, ... ]
      - "jds": { "<title>": { "mission": str, "requirements": [...], "nice_to_haves": [...], "benefits": [...] }, ... }
    """
    doc = Document()
    doc.add_heading("Hiring Plan", level=0)

    # --- Header: timeline + optional context (budget, location) ---
    weeks = plan.get("timeline_weeks")
    budget = plan.get("budget_usd")
    location = plan.get("location_policy")
    header_bits = []
    if weeks:
        header_bits.append(f"Target timeline: {weeks} week(s)")
    if budget is not None:
        header_bits.append(f"Budget: ${budget:,}")
    if location:
        header_bits.append(f"Location policy: {location}")
    if header_bits:
        doc.add_paragraph(" · ".join(header_bits))

    # --- Checklist (tasks) ---
    doc.add_heading("Checklist", level=1)
    for t in plan.get("tasks", []) or []:
        # Be defensive: pull with defaults to avoid KeyErrors
        name = str(t.get("name", "")).strip()
        owner = str(t.get("owner", "HR")).strip()
        due = str(t.get("due", "TBD")).strip()
        if not name:
            continue
        line = f"{name} — owner: {owner}, due: {due}"
        doc.add_paragraph(line, style="List Bullet")

    # --- Interview Loop ---
    doc.add_heading("Interview Loop", level=1)
    for s in plan.get("interview_loop", []) or []:
        stage = str(s.get("stage", "")).strip()
        if not stage:
            continue
        duration = s.get("duration_min", 45)
        try:
            duration = int(duration)
        except Exception:
            duration = 45
        signals = _norm_list(s.get("signals"))
        line = f"{stage} ({duration} min): {', '.join(signals) if signals else '—'}"
        doc.add_paragraph(line, style="List Bullet")

    # --- Roles & JDs ---
    doc.add_heading("Roles & JDs", level=1)
    jds = plan.get("jds", {}) or {}

    # Render each role JD as a subsection (H2). This is a summary—per-role DOCX can be exported separately.
    for title, jd in jds.items():
        doc.add_heading(str(title), level=2)
        mission = (jd.get("mission") or "").strip()
        if mission:
            doc.add_paragraph(f"Mission: {mission}")

        req = _norm_list(jd.get("requirements"))
        if req:
            doc.add_paragraph("Requirements:")
            _add_bullets(doc, req, style="List Bullet 2")

        nice = _norm_list(jd.get("nice_to_haves"))
        if nice:
            doc.add_paragraph("Nice-to-haves:")
            _add_bullets(doc, nice, style="List Bullet 2")

        benefits = _norm_list(jd.get("benefits"))
        if benefits:
            doc.add_paragraph("Benefits:")
            _add_bullets(doc, benefits, style="List Bullet 2")

    # --- Save to bytes ---
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ----------------------------------------------------------------------
# Per-role JD export
# ----------------------------------------------------------------------

def jd_to_docx(jd: Any) -> bytes:
    """
    Build a single-role JD .docx from either:
      - a Pydantic JD model (has .model_dump()), or
      - a plain dict with keys:
          title, mission, responsibilities, requirements, nice_to_haves, benefits

    Returns raw bytes for direct download.
    """
    # Accept pydantic model or plain dict
    data = jd.model_dump() if hasattr(jd, "model_dump") else dict(jd)

    title = (data.get("title") or "Job Description").strip()
    mission = (data.get("mission") or "").strip()
    responsibilities = _norm_list(data.get("responsibilities"))
    requirements = _norm_list(data.get("requirements"))
    nice_to_haves = _norm_list(data.get("nice_to_haves"))
    benefits = _norm_list(data.get("benefits"))

    doc = Document()
    doc.add_heading(f"Job Description — {title}", level=1)

    if mission:
        doc.add_heading("Mission", level=2)
        doc.add_paragraph(mission)

    if responsibilities:
        doc.add_heading("Responsibilities", level=2)
        _add_bullets(doc, responsibilities, style="List Bullet")

    if requirements:
        doc.add_heading("Requirements", level=2)
        _add_bullets(doc, requirements, style="List Bullet")

    if nice_to_haves:
        doc.add_heading("Nice-to-haves", level=2)
        _add_bullets(doc, nice_to_haves, style="List Bullet")

    if benefits:
        doc.add_heading("Benefits", level=2)
        _add_bullets(doc, benefits, style="List Bullet")

    # Small, non-intrusive footer to indicate fluidity of startup roles
    p = doc.add_paragraph()
    footer = p.add_run("\nThis description outlines typical duties and qualifications and may evolve.")
    footer.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ----------------------------------------------------------------------
# Bulk (all JDs) ZIP export
# ----------------------------------------------------------------------

def jds_to_zip(jds_map: Dict[str, Any]) -> bytes:
    """
    Bundle multiple JDs into a single ZIP.
    - Keys are role titles.
    - Values are JD dicts or models accepted by `jd_to_docx`.

    Returns raw bytes of the .zip file.
    """
    mem = BytesIO()
    with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as zf:
        for title, jd in (jds_map or {}).items():
            try:
                data = jd_to_docx(jd)
                fname = f"JD_{_safe_name(str(title))}.docx"
                zf.writestr(fname, data)
            except Exception:
                # Fail gracefully per-file; skip problematic entries and continue zipping others.
                continue
    mem.seek(0)
    return mem.getvalue()
